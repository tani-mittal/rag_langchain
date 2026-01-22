from docx import Document
from docx.document import Document as DocumentType
from typing import List
from pathlib import Path
from .base_processor import BaseDocumentProcessor, DocumentChunk

class DOCXProcessor(BaseDocumentProcessor):
    def extract_content(self, file_path: Path) -> List[DocumentChunk]:
        chunks = []
        
        try:
            self.logger.info(f"Processing DOCX: {file_path}")
            doc = Document(file_path)
            
            # Extract text by paragraphs and sections
            current_section = ""
            section_counter = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading
                if para.style.name.startswith('Heading'):
                    if current_section:
                        # Process previous section
                        chunks.extend(self._process_section(
                            current_section, 
                            file_path, 
                            section_counter,
                            len(chunks)
                        ))
                    current_section = text + "\n"
                    section_counter += 1
                else:
                    current_section += text + "\n"
            
            # Process last section
            if current_section:
                chunks.extend(self._process_section(
                    current_section, 
                    file_path, 
                    section_counter,
                    len(chunks)
                ))
            
            # If no sections found, process entire document
            if not chunks:
                full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                full_text = self.clean_text(full_text)
                
                if full_text:
                    chunks.extend(self._process_section(
                        full_text,
                        file_path,
                        1,
                        0
                    ))
            
            # Extract images
            images = self._extract_images(doc)
            if images and chunks:
                chunks[0].images = self.process_images(images)
            
            self.logger.info(f"Successfully processed DOCX: {len(chunks)} chunks created")
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_section(self, section_text: str, file_path: Path, section_num: int, chunk_offset: int) -> List[DocumentChunk]:
        section_text = self.clean_text(section_text)
        if not section_text:
            return []
            
        text_chunks = self.split_text(section_text)
        chunks = []
        
        for i, chunk_text in enumerate(text_chunks):
            chunk_id = self.create_chunk_id(str(file_path), chunk_offset + i)
            
            chunk = DocumentChunk(
                content=chunk_text,
                metadata={
                    "source": str(file_path),
                    "section": section_num,
                    "chunk_index": i,
                    "file_type": "docx",
                    "file_name": file_path.name
                },
                chunk_id=chunk_id,
                source_file=str(file_path)
            )
            chunks.append(chunk)
        
        return chunks
    
    def _extract_images(self, doc: DocumentType) -> List[bytes]:
        images = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        images.append(image_data)
                    except Exception as e:
                        self.logger.warning(f"Error extracting image: {e}")
                        continue
        except Exception as e:
            self.logger.warning(f"Error accessing document relationships: {e}")
        
        return images